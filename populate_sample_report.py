import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def populate_docx():
    template_path = r'D:\Projects\Mining Quality Prediction\Sample_InternshipReport_USC_UCT.docx'
    doc = Document(template_path)

    # 1. Update Title Page
    for p in doc.paragraphs:
        if 'Project Name' in p.text or '“Project Name”' in p.text or 'Project Name' in p.text:
            p.text = p.text.replace('“Project Name”', 'Quality Prediction in a Mining Process')
            p.text = p.text.replace('Project Name', 'Quality Prediction in a Mining Process')
            p.text = p.text.replace('Project Name', 'Quality Prediction in a Mining Process')
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        if '[Student name]' in p.text:
            p.text = p.text.replace('[Student name]', 'Vivaan Mahida')
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)

    # 2. Update Table 0 (Executive Summary)
    if len(doc.tables) > 0 and len(doc.tables[0].rows) > 1:
        cell = doc.tables[0].cell(1, 0)
        cell.text = (
            "This report provides details of the Industrial Internship provided by upskill Campus and The IoT Academy "
            "in collaboration with Industrial Partner UniConverge Technologies Pvt Ltd (UCT). This internship was focused on "
            "a project/problem statement provided by UCT. We had to finish the project including the report in 6 weeks’ time.\n\n"
            "My project was Quality Prediction in a Mining Process. The primary goal was to use high-frequency industrial IIoT sensor "
            "data from an iron ore flotation plant to predict % Silica Concentrate (impurity) in real time. Through temporal feature "
            "engineering (lags and rolling statistics) and advanced tree-based ensemble modeling (LightGBM, XGBoost, Random Forest), "
            "we achieved high minute-level prediction accuracy (R² = 0.6071, RMSE = 0.7196), validated a 1-hour early-warning forecast horizon, "
            "and proved that process control variables alone enable high-accuracy predictions without relying on manual laboratory iron measurements.\n\n"
            "This internship gave me a very good opportunity to get exposure to Industrial problems and design/implement solution for that. "
            "It was an overall great experience to have this internship."
        )

    # Helper function to find paragraph by start text
    def find_p_index(start_text):
        for idx, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(start_text):
                return idx
        return -1

    # Map of text replacements for main section guidance paragraphs
    replacements = {
        "Summary of the whole 6 weeks": (
            "Week-by-Week Project Summary:\n"
            "• Week 1 (Problem Onboarding & Domain Analysis): Studied froth flotation operations, physical process variables (airflow rates, column levels, pH, chemical reagents), and defined machine learning objectives.\n"
            "• Week 2 (Data Cleaning & Exploration): Processed 737,453 raw industrial sensor rows, handled sampling rate disparities (20s sensors vs 1h lab measurements), and analyzed feature correlations.\n"
            "• Week 3 (Temporal Feature Engineering): Built 1-minute and 1-hour uniform resampled datasets, engineered autoregressive lag features (1 to 60 minutes), rolling statistics (mean, std), and cyclical time encodings.\n"
            "• Week 4 (Model Development & Experiments): Trained XGBoost, LightGBM, and Random Forest regressors across three experiments: minute-level prediction, multi-step horizon analysis, and iron feature ablation.\n"
            "• Week 5 (Evaluation & Validation): Evaluated models using RMSE, MAE, and R² metrics, established 1-hour forecast reliability limits, and enforced chronological train/test splits.\n"
            "• Week 6 (Dashboard & Reporting): Built a standalone interactive Plotly HTML dashboard, finalized technical documentation, and published the open-source codebase to GitHub."
        ),
        "About need of relevant Internship": (
            "Relevance of Internship in Career Development:\n"
            "Industrial internships are essential for bridging academic theory and enterprise application. Working with real-world sensor data "
            "under production constraints provides crucial exposure to software design patterns, time-series validation protocols, and industrial problem solving."
        ),
        "Brief about Your project": (
            "Project Overview:\n"
            "The project focuses on predicting product quality (% Silica Concentrate) in an iron ore flotation plant. By replacing periodic 1-hour lab delays "
            "with minute-level machine learning predictions, plant engineers receive an early warning to adjust process parameters, saving valuable ore and reducing waste."
        ),
        "Opportunity given by USC/UCT": (
            "Opportunity Provided by USC & UCT:\n"
            "I extend my sincere appreciation to upskill Campus, The IoT Academy, and UniConverge Technologies Pvt Ltd for facilitating this hands-on industrial internship."
        ),
        "How Program was planned": (
            "Program Structure:\n"
            "The program was structured into weekly milestones covering domain orientation, exploratory data analysis, feature engineering, model optimization, dashboard creation, and final reporting."
        ),
        "Your Learnings and overall experience": (
            "Learnings & Overall Experience:\n"
            "Gained mastery in time-series feature engineering, gradient boosting (LightGBM/XGBoost), Plotly visualization, and industrial IIoT process control."
        ),
        "Thank to all (with names)": (
            "Acknowledgments:\n"
            "Special thanks to the technical mentors and coordinators at UniConverge Technologies Pvt Ltd, upskill Campus, and The IoT Academy for their continuous guidance and feedback."
        ),
        "Your message to your juniors": (
            "Message to Juniors & Peers:\n"
            "Always take time to understand the physical domain behind industrial data, enforce strict validation to avoid temporal data leakage, and present results in business-ready visual formats."
        ),
    }

    for p in doc.paragraphs:
        for key, new_txt in replacements.items():
            if key in p.text:
                p.text = new_txt
                break

    # Glossary Table 1 Population
    if len(doc.tables) > 1:
        g_table = doc.tables[1]
        g_data = [
            ("% Silica Concentrate", "Target impurity variable in iron ore concentrate; lower values indicate higher ore quality."),
            ("Flotation Plant", "Industrial facility using chemical reagents and air bubbles to separate iron minerals from silica sand."),
            ("Tailings", "Waste product discharged into environmental ponds when iron ore is lost during processing."),
            ("Lag Feature", "Time-series feature derived from historical past values (e.g., target value at t-5 minutes)."),
            ("RMSE / MAE / R²", "Statistical metrics measuring Root Mean Squared Error, Mean Absolute Error, and Variance Explained."),
        ]
        # Ensure enough rows
        while len(g_table.rows) - 1 < len(g_data):
            g_table.add_row()
        for idx, (term, defn) in enumerate(g_data, start=1):
            g_table.cell(idx, 0).text = term
            g_table.cell(idx, 1).text = defn
            g_table.cell(idx, 0).paragraphs[0].runs[0].bold = True

    # Problem Statement Section
    p_prob_idx = find_p_index("In the assigned problem statement")
    if p_prob_idx != -1:
        doc.paragraphs[p_prob_idx].text = (
            "Problem Statement & Background:\n"
            "In iron ore flotation plants, reagents (Starch and Amina) and air bubbles are injected into flotation columns to float valuable iron minerals "
            "while allowing silica (sand) impurities to sink as waste tailings. The primary quality goal is to minimize the percentage of Silica Concentrate (% Silica Concentrate) in the final pulp.\n\n"
            "However, measuring % Silica Concentrate currently requires physical sampling and chemical lab analysis, which is performed ONLY ONCE PER HOUR. "
            "This creates a 59-minute operational blind spot. If process fluctuations or chemical imbalances cause a silica spike, engineers discover it too late, "
            "resulting in sub-standard ore output or valuable iron being lost to environmental tailings ponds.\n\n"
            "Core Research Questions Addressed:\n"
            "1. Is it possible to predict % Silica Concentrate every minute using continuous IIoT sensor data?\n"
            "2. How many hours ahead can we reliably forecast % Silica Concentrate to enable proactive operational control?\n"
            "3. Is it possible to predict % Silica Concentrate without using the highly correlated % Iron Concentrate column (enabling true lab-free real-time operation)?"
        )
        if p_prob_idx + 1 < len(doc.paragraphs) and "[Explain your problem statement]" in doc.paragraphs[p_prob_idx + 1].text:
            doc.paragraphs[p_prob_idx + 1].text = ""

    # Existing and Proposed Solution
    p_sol_idx = find_p_index("Provide summary of existing solutions")
    if p_sol_idx != -1:
        doc.paragraphs[p_sol_idx].text = (
            "Limitations of Existing Solutions:\n"
            "• Manual Laboratory Testing: Operations rely on periodic 1-hour lab testing, introducing 60 minutes of feedback latency.\n"
            "• Reactive Control: Operators react after quality degradation has already occurred.\n"
            "• Linear Rule Systems: Traditional control systems use static linear rules that fail to capture non-linear interactions across 7 flotation columns.\n\n"
            "Proposed Machine Learning Solution:\n"
            "We developed an automated time-series ML pipeline that ingests continuous 20-second sensor readings (airflows, column levels, pH, chemical flows), "
            "resamples data into 1-minute intervals, generates temporal lag and rolling features, and predicts % Silica Concentrate using LightGBM and XGBoost.\n\n"
            "Value Addition & Impact:\n"
            "• Minute-Level Quality Visibility: Real-time quality predictions every minute.\n"
            "• 1-Hour Early Warning: 60-minute advance notice to adjust chemical dosing and airflows.\n"
            "• Lab Independence: Accurate predictions without requiring lab-tested % Iron Concentrate.\n"
            "• Waste & Cost Reduction: Minimizes iron ore loss to tailings, saving costs and protecting the environment."
        )
        for offset in [1, 2]:
            if p_sol_idx + offset < len(doc.paragraphs):
                if any(k in doc.paragraphs[p_sol_idx + offset].text for k in ["What is your proposed", "What value addition"]):
                    doc.paragraphs[p_sol_idx + offset].text = ""

    # Code and Report Submission links
    for p in doc.paragraphs:
        if "Code submission (Github link)" in p.text:
            p.text = "Code submission (Github link):\nhttps://github.com/vivaanmahida/mining-quality-prediction"
        if "Report submission (Github link)" in p.text:
            p.text = "Report submission (Github link):\nhttps://github.com/vivaanmahida/mining-quality-prediction/blob/main/Internship_Report_Mining_Quality_Prediction.docx"

    # Proposed Design / Model Section
    p_des_idx = find_p_index("Given more details about design flow")
    if p_des_idx != -1:
        doc.paragraphs[p_des_idx].text = (
            "System Design & Architecture:\n"
            "The solution follows a modular 4-stage data engineering and machine learning pipeline:\n\n"
            "1. Data Preprocessor (01_eda.py): Cleans sensor data, handles timestamps, and resamples 20s data into 1-minute and 1-hour uniform grids.\n"
            "2. Feature Engineer (02_feature_engineering.py): Computes autoregressive lag features (target_lag_1 to target_lag_60), rolling window statistics (5m to 60m mean/std), cyclical time encodings, and multi-step target horizons (1h to 12h ahead).\n"
            "3. Model Trainer (03_train_models.py): Executes chronological train/test splits (80% train, 20% test) and trains XGBoost, LightGBM, and Random Forest regressors across 3 experiments.\n"
            "4. Dashboard Generator (04_generate_dashboard.py): Encapsulates model outputs, feature importances, and predictions into a standalone Plotly HTML dashboard (dashboard/index.html)."
        )

    # Performance Test Section
    p_perf_idx = find_p_index("This is very important part and defines why this work")
    if p_perf_idx != -1:
        doc.paragraphs[p_perf_idx].text = (
            "Constraints & Performance Evaluation:\n\n"
            "Key Constraints Addressed:\n"
            "• Temporal Data Leakage: Evaluated models using strict chronological time splits (80% train, 20% test) rather than random cross-validation.\n"
            "• Edge Latency: Inference executes in milliseconds per sample, making it suitable for edge deployment on IIoT gateways.\n"
            "• Accuracy Benchmark: Target R² > 0.50 for minute-level prediction and RMSE below baseline standard deviation (1.12%).\n\n"
            "Experimental Performance Outcomes:\n\n"
            "1. Experiment A — Minute-Level Quality Prediction (Q1):\n"
            "   - LightGBM (Best Model): RMSE = 0.7196 | MAE = 0.5301 | R² = 0.6071 ✅\n"
            "   - XGBoost Regressor: RMSE = 0.7534 | MAE = 0.5760 | R² = 0.5692\n"
            "   - Random Forest: RMSE = 0.7775 | MAE = 0.5995 | R² = 0.5412\n\n"
            "2. Experiment B — Multi-Step Forecast Horizon Analysis (Q2):\n"
            "   - 1 Hour Ahead: RMSE = 0.8467 | MAE = 0.6503 | R² = 0.4558 (Operational / Usable ⚠️)\n"
            "   - 2 Hours Ahead: RMSE = 0.9435 | MAE = 0.7406 | R² = 0.3242 (Weak Signal)\n"
            "   - 4 Hours Ahead: RMSE = 1.0809 | MAE = 0.8809 | R² = 0.1133 (Unreliable ❌)\n\n"
            "3. Experiment C — Iron Feature Ablation Study (Q3):\n"
            "   - XGBoost WITH % Iron Concentrate: RMSE = 0.7534 | R² = 0.5692\n"
            "   - XGBoost WITHOUT % Iron Concentrate: RMSE = 0.7548 | R² = 0.5677 ✅\n"
            "   - Key Finding: Removing % Iron Concentrate results in negligible accuracy loss (ΔRMSE = +0.0014), proving true lab independence!"
        )
        for offset in range(1, 6):
            if p_perf_idx + offset < len(doc.paragraphs):
                if any(k in doc.paragraphs[p_perf_idx + offset].text for k in ["Here we need to first find", "How those constraints", "What were test results", "Constraints can be", "In case you could not"]):
                    doc.paragraphs[p_perf_idx + offset].text = ""

    # My Learnings Section
    p_learn_idx = find_p_index("You should provide summary of your overall learning")
    if p_learn_idx != -1:
        doc.paragraphs[p_learn_idx].text = (
            "Summary of Overall Learnings:\n"
            "• Time-Series ML: Hands-on experience structuring high-frequency sensor data, engineering lag features, and enforcing chronological validation.\n"
            "• Gradient Boosting: Tuning hyperparameter bounds and early stopping for LightGBM and XGBoost on industrial tabular data.\n"
            "• Domain Knowledge: Deep understanding of froth flotation plant dynamics and IIoT process control.\n"
            "• Full-Stack Visualization: Creating single-file interactive Plotly HTML dashboards for enterprise stakeholders.\n"
            "• Professional Standards: Git/GitHub version control workflows and comprehensive technical documentation."
        )

    # Future Work Scope Section
    p_fut_idx = find_p_index("You can put some ideas that you could not work")
    if p_fut_idx != -1:
        doc.paragraphs[p_fut_idx].text = (
            "Future Work & Potential Scope:\n"
            "1. Deep Learning Sequence Models: Explore Temporal Convolutional Networks (TCN) or LSTM architectures for multi-hour forecasting.\n"
            "2. UCT Smart Factory Integration: Deploy trained LightGBM models to UCT Insight platform via MQTT/REST endpoints for live edge inference.\n"
            "3. Automated Tuning: Integrate Optuna for automated continuous hyperparameter tuning as sensor calibration drifts.\n"
            "4. Prescriptive Control Feedback: Extend predictive models to prescribe optimal column airflow and chemical dosing setpoints."
        )

    # Save output docx
    out_file1 = r'D:\Projects\Mining Quality Prediction\Internship_Report_Mining_Quality_Prediction.docx'
    out_file2 = r'D:\Projects\Mining Quality Prediction\mining-quality-prediction\Internship_Report_Mining_Quality_Prediction.docx'
    
    doc.save(out_file1)
    doc.save(out_file2)
    print("Populated template successfully saved to:", out_file1)

if __name__ == '__main__':
    populate_docx()
