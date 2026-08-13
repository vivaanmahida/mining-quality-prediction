import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client

def create_submission_files():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    NAVY = RGBColor(0x1B, 0x36, 0x5D)
    BLUE = RGBColor(0x00, 0x66, 0xCC)
    DARK = RGBColor(0x22, 0x22, 0x22)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = BLUE
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        return p

    # Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("ONLINE INTERNSHIP TESTIMONIAL SUBMISSION")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(20)
    r_t.font.bold = True
    r_t.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("upSkill Campus & UniConverge Technologies (P) Ltd Internship Program")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = BLUE

    # Learner Metadata Box Table
    table_meta = doc.add_table(rows=6, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Learner Name:", "Vivaan Mahida"),
        ("Domain / Program:", "Data Science & Machine Learning (6-Week Industrial Internship)"),
        ("Project Name:", "Quality Prediction in a Mining Process"),
        ("Facilitators:", "upSkill Campus & The IoT Academy"),
        ("Industrial Partner:", "UniConverge Technologies (P) Ltd"),
        ("GitHub Repository:", "https://github.com/vivaanmahida/upskillcampus"),
    ]
    for idx, (label, val) in enumerate(meta_data):
        c0 = table_meta.cell(idx, 0)
        c1 = table_meta.cell(idx, 1)
        c0.paragraphs[0].text = label
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = NAVY
        c1.paragraphs[0].text = val
        if label.startswith("GitHub"):
            c1.paragraphs[0].runs[0].font.color.rgb = BLUE
            c1.paragraphs[0].runs[0].bold = True

    doc.add_paragraph() # spacing

    # Section 1: Video Testimonial Script
    add_h1("1. Video Testimonial Script (Spoken Transcript)")
    add_p("Below is the exact transcript recorded for the online internship video testimonial based on the official format provided in 'Online internship testimonial format.mov':")
    
    p_script = doc.add_paragraph()
    p_script.paragraph_format.space_before = Pt(6)
    p_script.paragraph_format.space_after = Pt(12)
    r_sc = p_script.add_run(
        "“Hello everyone! My name is Vivaan Mahida, and I have recently completed a 6-week project-based Industrial Internship "
        "in Data Science & Machine Learning from upSkill Campus in collaboration with UniConverge Technologies (P) Ltd.\n\n"
        "During this internship, I worked on a real-world industrial project titled 'Quality Prediction in a Mining Process'. "
        "My main objective was to build Machine Learning models to predict silica impurity (% Silica Concentrate) in an iron ore "
        "flotation plant using real IIoT sensor data. Through this project, I engineered time-series lag features, trained advanced "
        "gradient boosting models like LightGBM and XGBoost achieving an R² of 0.607, validated a 1-hour early-warning forecast horizon, "
        "and developed an interactive Plotly dashboard.\n\n"
        "This program is truly more than just an online internship because I gained real practical knowledge, enterprise-level coding experience, "
        "and hands-on exposure to industrial IoT problem-solving.\n\n"
        "I would like to express my heartfelt thanks to upSkill Campus, The IoT Academy, and UniConverge Technologies (P) Ltd for this "
        "incredible opportunity and their constant support throughout the journey. If you are looking to enhance your practical skills, "
        "I highly recommend signing up on their website upskillcampus.com. Thank you!”"
    )
    r_sc.font.italic = True
    r_sc.font.size = Pt(11)
    r_sc.font.color.rgb = DARK

    # Section 2: Social Media Post & Tagging Verification
    add_h1("2. Social Media Post & Tagging Verification")
    add_p("Per the assignment instructions, this video testimonial and project experience are shared on social media (LinkedIn) with mandatory partner tagging:")
    
    add_h2("LinkedIn Post Text:")
    p_post = doc.add_paragraph()
    r_po = p_post.add_run(
        "🚀 Excited to share that I have successfully completed my 6-week Industrial Internship in Data Science & Machine Learning "
        "facilitated by @upSkill Campus and @The IoT Academy in collaboration with industrial partner @UniConverge Technologies (P) Ltd!\n\n"
        "📌 Project Title: Quality Prediction in a Mining Process\n"
        "🔍 Implementation: Engineered time-series lag features on 737,453 sensor readings across 183 days, benchmarked LightGBM & XGBoost "
        "achieving R² = 0.6071, validated a 1-hour early-warning forecast horizon, and built an interactive Plotly HTML dashboard.\n\n"
        "💻 GitHub Repo: https://github.com/vivaanmahida/upskillcampus\n"
        "📄 Official Report: https://github.com/vivaanmahida/upskillcampus/blob/main/QualityPredictionInAMiningProcess_Vivaan_USC_UCT.pdf\n\n"
        "Special thanks to upSkill Campus and UniConverge Technologies (P) Ltd for this practical industrial learning experience!\n"
        "#DataScience #MachineLearning #upSkillCampus #UniConvergeTechnologies #IIoT #Python"
    )
    r_po.font.size = Pt(10.5)

    # Section 3: Official Deliverable Submission Links
    add_h1("3. Official Project Submission Links")
    add_p("1. GitHub Repository: https://github.com/vivaanmahida/upskillcampus")
    add_p("2. Code Submission File: https://github.com/vivaanmahida/upskillcampus/blob/main/QualityPredictionInAMiningProcess.py")
    add_p("3. Project Report PDF: https://github.com/vivaanmahida/upskillcampus/blob/main/QualityPredictionInAMiningProcess_Vivaan_USC_UCT.pdf")

    # Save DOCX & PDF
    docx_filename = "Vivaan_Mahida_Internship_Testimonial_Submission.docx"
    pdf_filename  = "Vivaan_Mahida_Internship_Testimonial_Submission.pdf"
    txt_filename  = "Vivaan_Mahida_Internship_Testimonial_Submission.txt"

    paths_docx = [
        os.path.join(r'D:\Internship\Uniconverge', docx_filename),
        os.path.join(r'D:\Projects\Mining Quality Prediction\mining-quality-prediction', docx_filename),
    ]
    paths_pdf = [
        os.path.join(r'D:\Internship\Uniconverge', pdf_filename),
        os.path.join(r'D:\Projects\Mining Quality Prediction\mining-quality-prediction', pdf_filename),
    ]

    for p_d in paths_docx:
        doc.save(p_d)
        print("Saved DOCX:", p_d)

    # Convert DOCX to PDF using Word
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc_word = word.Documents.Open(os.path.abspath(paths_docx[1]))
    doc_word.SaveAs(os.path.abspath(paths_pdf[1]), FileFormat=17) # PDF
    doc_word.Close()
    word.Quit()

    with open(paths_pdf[1], 'rb') as f_in, open(paths_pdf[0], 'wb') as f_out:
        f_out.write(f_in.read())

    print("Saved PDF:", paths_pdf[0])

if __name__ == '__main__':
    create_submission_files()
